"""
VideoRipper — Flask backend.
Auth is handled client-side by hub-nav.js; server routes are unauthenticated.
"""
import io
import json
import os
import uuid
import threading
from datetime import datetime, timezone
from pathlib import Path

from flask import Flask, jsonify, request, send_from_directory, abort, Response

import ripper
import rev_client
import gdrive

# ── paths ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent
FRONTEND_DIR = BASE_DIR / 'frontend'
# DATA_DIR: use /data (Railway volume mount point) when it exists, else fall
# back to the repo-relative data/ dir for local dev.
_volume = Path('/data')
DATA_DIR = _volume if _volume.is_dir() else BASE_DIR / 'data'
VIDEOS_DIR = DATA_DIR / 'videos'
SCREENSHOTS_DIR = DATA_DIR / 'screenshots'
MANIFEST_PATH = DATA_DIR / 'manifest.json'

VIDEOS_DIR.mkdir(parents=True, exist_ok=True)
SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)

# ── Flask app ───────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder=str(FRONTEND_DIR))
# The bookmarklet posts hero-image data URLs captured in the user's browser
# (the only place a Cloudflare-gated promo page renders), so allow a large body.
app.config['MAX_CONTENT_LENGTH'] = 30 * 1024 * 1024


@app.after_request
def _add_cors_headers(resp):
    # The bookmarklet runs on the promo page's origin and posts here cross-origin.
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS, DELETE'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return resp


_manifest_lock = threading.Lock()


# ── manifest helpers ────────────────────────────────────────────────────────

_MANIFEST_BAK = MANIFEST_PATH.with_name('manifest.json.bak')
_MANIFEST_TMP = MANIFEST_PATH.with_name('manifest.json.tmp')


def _load_manifest() -> list:
    if not MANIFEST_PATH.exists():
        return []
    try:
        with open(MANIFEST_PATH) as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        # The live file is corrupt/partial (e.g. a restart interrupted a write).
        # Recover from the last-good backup. Crucially, do NOT fall back to an
        # empty list — a writer would then persist [] and destroy every job.
        if _MANIFEST_BAK.exists():
            try:
                with open(_MANIFEST_BAK) as f:
                    return json.load(f)
            except (json.JSONDecodeError, OSError):
                pass
        raise RuntimeError('manifest.json is unreadable and no valid backup exists')


def _save_manifest(jobs: list) -> None:
    # Atomic write: dump to a temp file in the same dir, fsync, back up the
    # current good file, then os.replace() (atomic on the same filesystem).
    # This guarantees the live manifest is never left truncated by a crash or
    # restart mid-write — the failure mode that was silently emptying it.
    with open(_MANIFEST_TMP, 'w') as f:
        json.dump(jobs, f, indent=2)
        f.flush()
        os.fsync(f.fileno())
    if MANIFEST_PATH.exists():
        try:
            import shutil
            shutil.copy2(MANIFEST_PATH, _MANIFEST_BAK)
        except OSError:
            pass
    os.replace(_MANIFEST_TMP, MANIFEST_PATH)


def _get_job(job_id: str) -> dict | None:
    jobs = _load_manifest()
    return next((j for j in jobs if j['id'] == job_id), None)


def _update_job(job_id: str, updates: dict) -> dict | None:
    with _manifest_lock:
        jobs = _load_manifest()
        for job in jobs:
            if job['id'] == job_id:
                job.update(updates)
                _save_manifest(jobs)
                return job
    return None


def _append_job(job: dict) -> None:
    with _manifest_lock:
        jobs = _load_manifest()
        jobs.append(job)
        _save_manifest(jobs)


def _delete_job(job_id: str) -> bool:
    with _manifest_lock:
        jobs = _load_manifest()
        new_jobs = [j for j in jobs if j['id'] != job_id]
        if len(new_jobs) == len(jobs):
            return False
        _save_manifest(new_jobs)
    return True


def _find_duplicate(source_url: str) -> dict | None:
    """Return the first completed job with the same source_url, or None."""
    jobs = _load_manifest()
    for job in jobs:
        if (job.get('source_url') == source_url
                and job.get('rev_status') == 'complete'
                and job.get('transcript_text')):
            return job
    return None


# ── Google Drive archival ────────────────────────────────────────────────────
# After ARCHIVE_AFTER_DAYS, upload each local video to the Shared Drive, record
# the Drive id/url on the job, and delete the local .mp4 to reclaim volume space.
# Playback keeps working because /video streams from Drive when local is gone.
ARCHIVE_AFTER_DAYS = 30


def _archive_old_videos(days: int | None = None) -> dict:
    """Sweep: archive local videos older than the retention window to Drive.
    Best-effort per job — a failure records archive_error and keeps the local
    file for the next sweep. No-ops (safely) until Drive is configured."""
    if not gdrive.is_configured():
        return {'skipped': 'google drive not configured'}
    cutoff_days = ARCHIVE_AFTER_DAYS if days is None else days
    now = datetime.now(timezone.utc)
    archived = errors = 0
    for job in _load_manifest():
        jid = job.get('id')
        if not jid or job.get('drive_file_id'):
            continue  # already archived
        vp = VIDEOS_DIR / f'{jid}.mp4'
        if not vp.exists():
            continue
        ts = job.get('created_at') or job.get('rev_submitted_at') or ''
        try:
            created = datetime.fromisoformat(ts)
        except ValueError:
            continue
        if (now - created).total_seconds() < cutoff_days * 86400:
            continue
        safe = ''.join(
            c for c in (job.get('title') or jid)[:60]
            if c.isalnum() or c in ' -_'
        ).strip() or jid
        try:
            up = gdrive.upload_video(str(vp), f'{safe} [{jid}].mp4')
        except Exception as exc:
            _update_job(jid, {'archive_error': str(exc)[:200]})
            errors += 1
            continue
        _update_job(jid, {
            'drive_file_id': up['file_id'],
            'drive_url': up.get('web_view_link', ''),
            'archived_at': now.isoformat(),
            'local_video': '',
            'archive_error': '',
        })
        try:
            vp.unlink()
        except OSError:
            pass
        archived += 1
    return {'archived': archived, 'errors': errors, 'cutoff_days': cutoff_days}


_archiver_started = False


def start_archiver() -> None:
    """Start the daily archival sweep in a daemon thread (idempotent)."""
    global _archiver_started
    if _archiver_started:
        return
    _archiver_started = True

    def _loop():
        import time
        time.sleep(120)  # let the app settle after boot
        while True:
            try:
                result = _archive_old_videos()
                print(f'[archiver] sweep: {result}', flush=True)
            except Exception as exc:
                print(f'[archiver] sweep error: {exc}', flush=True)
            time.sleep(24 * 3600)

    threading.Thread(target=_loop, daemon=True).start()


# ── docx generation ─────────────────────────────────────────────────────────

def _build_docx(title: str, transcript_text: str, eyebrow: str = '',
                headline: str = '', subhead: str = '', subhead2: str = '') -> bytes:
    """Return .docx bytes for the given transcript.

    When a headline block was extracted from the promo page, render it above the
    transcript in descending prominence — eyebrow (small), headline (largest),
    subhead, then a secondary subhead — all bold. Falls back to the plain title
    heading when no headline block is available.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        # (text, point size) in visual order; skipped when empty.
        header_lines = [
            (eyebrow, 13),
            (headline, 26),
            (subhead, 16),
            (subhead2, 14),
        ]
        if any(text for text, _ in header_lines):
            for text, size in header_lines:
                if not text:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(size)
            doc.add_paragraph()  # spacer before the transcript
        else:
            doc.add_heading(title or 'Transcript', level=1)
        # Split into paragraphs on double newline, add each as a paragraph
        paragraphs = [p.strip() for p in transcript_text.split('\n\n') if p.strip()]
        if not paragraphs:
            paragraphs = [transcript_text]
        for para in paragraphs:
            doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        raise RuntimeError('python-docx is not installed')


# ── title composition ─────────────────────────────────────────────────────

def _compose_title(source_url: str, video_title: str, page_title: str) -> str:
    """
    Build a friendly job title: "{Publisher} - {video title}".
    Falls back to the page title, then to whatever is available.
    """
    publisher = ripper.publisher_from_url(source_url)
    vt = (video_title or '').strip()
    pt = (page_title or '').strip()
    # Don't use a page "title" that is really just the source URL.
    if pt == source_url:
        pt = ''
    chosen = vt or pt
    if publisher and chosen:
        return f'{publisher} - {chosen}'
    return chosen or publisher or (pt or source_url)


def _apply_video_metadata(job_id: str, source_url: str, platform: str, video_id: str,
                          extra: dict, page_title: str = '') -> None:
    """
    Probe yt-dlp for the real video title + thumbnail and update the job.
    Best-effort: never raises. Sets title to "{Publisher} - {video title}" and,
    when the page had no og:image thumbnail, backfills info['thumbnail'].
    """
    info = ripper.probe_video_info(platform, video_id, source_url, extra)
    video_title = ripper.video_title_from_info(info)
    updates = {'title': _compose_title(source_url, video_title, page_title)}
    job = _get_job(job_id)
    if (not job or not job.get('thumbnail_url')) and info.get('thumbnail'):
        updates['thumbnail_url'] = info['thumbnail']
    _update_job(job_id, updates)


# ── background pipeline ─────────────────────────────────────────────────────

def _take_screenshot(job_id: str, target_url: str) -> None:
    """Full-page screenshot of target_url; save to SCREENSHOTS_DIR/{job_id}.png.

    Captures the ENTIRE scrollable page (scroll + stitch, like GoFullPage) so
    the promo headline and full sales copy are preserved. Renders direct first;
    if the page comes back as a Cloudflare bot-challenge or a blank/broken render
    (datacenter IPs get challenged, e.g. Porter), retries through the residential
    proxy — a trusted IP that renders the real page — when one is configured
    (CNN_PROXY / RESIDENTIAL_PROXY_URL). Loads any uploaded cookies too.
    """
    proxy = ripper._playwright_proxy(ripper.cnn_proxy_url())
    try:
        blocked = _capture_page(job_id, target_url, None)
        if blocked and proxy:
            _capture_page(job_id, target_url, proxy)
        _update_job(job_id, {'has_screenshot': True})
    except Exception as exc:
        # Non-fatal — screenshot failure should not block transcription
        _update_job(job_id, {'screenshot_error': str(exc)[:200]})


def _capture_page(job_id: str, target_url: str, proxy) -> bool:
    """Render target_url (optionally via a proxy) and save the full-page +
    top-clip screenshots. Returns True if the render looks like a Cloudflare
    challenge or a blank/broken page (so the caller can retry through a proxy)."""
    blocked = False
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True,
            proxy=proxy,
            args=['--no-sandbox', '--disable-setuid-sandbox'],
        )
        context = browser.new_context(
            user_agent=(
                'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/120.0.0.0 Safari/537.36'
            ),
            viewport={'width': 1280, 'height': 800},
        )
        pw_cookies = ripper._load_cookies_for_playwright(target_url)
        if pw_cookies:
            context.add_cookies(pw_cookies)
        page = context.new_page()
        try:
            from playwright_stealth import stealth_sync
            stealth_sync(page)
        except ImportError:
            pass
        if True:
            try:
                page.goto(target_url, wait_until='domcontentloaded', timeout=30000)
                page.wait_for_load_state('networkidle', timeout=8000)
            except Exception:
                pass
            # Detect a Cloudflare bot-challenge or a blank/broken render so the
            # caller can retry via a trusted (residential-proxy) IP.
            try:
                sig = page.evaluate("""() => {
                    const t = (document.body ? document.body.innerText : '').trim().length;
                    let media = 0;
                    document.querySelectorAll('video,iframe,img').forEach((e) => {
                        const r = e.getBoundingClientRect();
                        if (r.width > 200 && r.height > 150) media++;
                    });
                    const html = document.documentElement.innerHTML.slice(0, 4000).toLowerCase();
                    const cf = /just a moment|verify you are human|cf-chl|challenge-platform|cdn-cgi\\/challenge|checking your browser/.test(html);
                    return { t, media, cf };
                }""")
                blocked = bool(sig.get('cf')) or (
                    int(sig.get('t') or 0) < 60 and int(sig.get('media') or 0) == 0
                )
            except Exception:
                pass
            # Long sales pages defer most images/sections until scrolled into
            # view. Step-scroll to the bottom to trigger lazy loading, then
            # return to the top before capturing the full page.
            try:
                page.evaluate("""async () => {
                    await new Promise((resolve) => {
                        let total = 0;
                        const step = 800;
                        const timer = setInterval(() => {
                            window.scrollBy(0, step);
                            total += step;
                            if (total >= document.body.scrollHeight) {
                                clearInterval(timer);
                                window.scrollTo(0, 0);
                                resolve();
                            }
                        }, 120);
                    });
                }""")
                page.wait_for_timeout(1500)
            except Exception:
                pass
            # Strip ONLY genuine popups/modals/consent gates so the screenshot
            # shows the real hero — never remove content merely for being fixed
            # or large (a promo hero is often a fixed, full-viewport container,
            # e.g. Porter/Brownstone's `div.above-fold`; removing it blanks the
            # page). Two high-confidence signals: (a) class/id/role tagged as a
            # modal/popup/consent/exit-intent widget, or (b) a very-high-z
            # (>=1000) fixed full-screen layer (a true modal backdrop). Low
            # z-index fixed heroes are left intact.
            try:
                page.evaluate("""() => {
                    const RX = /(modal|popup|pop-up|overlay|lightbox|backdrop|consent|cookie|gdpr|interstitial|dialog|exit-intent|exitintent|ouibounce|optin|opt-in|klaviyo|privy|leadbox|fancybox)/i;
                    const kill = [];
                    document.querySelectorAll('body *').forEach((el) => {
                        const s = getComputedStyle(el);
                        if (s.display === 'none' || s.visibility === 'hidden') return;
                        const pos = s.position;
                        if (pos !== 'fixed' && pos !== 'absolute') return;
                        const z = parseInt(s.zIndex) || 0;
                        const r = el.getBoundingClientRect();
                        const tag = ((el.className || '') + ' ' + (el.id || '')).toString();
                        const tagged = RX.test(tag)
                            || el.getAttribute('role') === 'dialog'
                            || el.getAttribute('aria-modal') === 'true';
                        const bigBackdrop = pos === 'fixed' && z >= 1000
                            && r.width >= window.innerWidth * 0.6
                            && r.height >= window.innerHeight * 0.5;
                        if ((tagged && z >= 1) || bigBackdrop) kill.push(el);
                    });
                    kill.forEach((el) => { try { el.remove(); } catch (e) {} });
                    for (const t of [document.documentElement, document.body]) {
                        if (t) t.style.overflow = 'auto';  // unlock modal scroll-lock
                    }
                }""")
                page.wait_for_timeout(300)
            except Exception:
                pass
            dest = str(SCREENSHOTS_DIR / f'{job_id}.png')
            page.screenshot(path=dest, full_page=True)
            # Also capture a top-of-page clip for the analyzer. The vision API
            # rejects images taller than 8000px and sales pages are far taller,
            # so send the top — where the headline/subheadline, the video and
            # its CTA button live — as a legible, in-limits crop.
            try:
                metrics = page.evaluate("""() => {
                    const sels = ['video','iframe','[class*="wistia"]','[data-video-id]','.video-js'];
                    let bottom = 0;
                    for (const sel of sels) {
                        for (const el of document.querySelectorAll(sel)) {
                            const r = el.getBoundingClientRect();
                            if (r.width > 200 && r.height > 100) {
                                const b = r.top + window.scrollY + r.height;
                                if (b > bottom) bottom = b;
                            }
                        }
                    }
                    return {
                        videoBottom: Math.ceil(bottom),
                        pageHeight: Math.ceil(document.documentElement.scrollHeight),
                        pageWidth: Math.ceil(document.documentElement.clientWidth) || 1280,
                    };
                }""")
                page_h = max(1, int(metrics.get('pageHeight') or 0))
                page_w = min(1280, max(1, int(metrics.get('pageWidth') or 1280)))
                v_bottom = int(metrics.get('videoBottom') or 0)
                # Include the video + ~500px below for the CTA button; fall back
                # to the first 3000px when no video is found.
                clip_h = (v_bottom + 500) if v_bottom > 0 else 3000
                clip_h = max(1, min(clip_h, 8000, page_h))
                top_dest = str(SCREENSHOTS_DIR / f'{job_id}_top.png')
                page.screenshot(
                    path=top_dest,
                    clip={'x': 0, 'y': 0, 'width': page_w, 'height': clip_h},
                )
            except Exception:
                pass  # top clip is best-effort; full-page PNG is still saved
            browser.close()
    return blocked


# Cheap, fast vision model for headline/subheadline extraction.
HEADLINE_MODEL = 'claude-haiku-4-5-20251001'


def _ensure_poster(job_id: str, platform: str = '', account_id: str = '',
                   embed_id: str = '') -> bool:
    """Make sure a poster image exists at SCREENSHOTS_DIR/{job_id}_poster.png.

    On gated promos the headline is baked into the video, and the promo page
    can't be screenshotted server-side. The headline is almost always the
    video's OPENING frame (its title card, e.g. Porter's "TRUMP'S NEW DOLLAR").
    So prefer an ffmpeg still from the start of the downloaded MP4 — the
    platform-native thumbnail (Vidalytics loader config) is a MID-video frame
    that misses the headline, so it's only a fallback when we have no video.
    """
    if list(SCREENSHOTS_DIR.glob(f'{job_id}_poster*.png')):
        return True
    # 1) Preferred: still frames from the START of the video. Grab several
    # early timestamps because a VSL's title card (the headline) can be brief
    # and t=0 may be a black fade-in; extraction leaves them for the vision
    # loop to try earliest-first.
    mp4 = VIDEOS_DIR / f'{job_id}.mp4'
    made = False
    if mp4.exists():
        for i, seek in enumerate((0, 0.5, 1.5, 3)):
            dest = SCREENSHOTS_DIR / f'{job_id}_poster{i}.png'
            try:
                if ripper.extract_poster_frame(str(mp4), str(dest), seek):
                    made = True
            except Exception:
                pass
        if made:
            return True
    # 2) Fallback (no local video): Vidalytics native thumbnail from the loader.
    if platform == 'vidalytics' and account_id and embed_id:
        try:
            url = ripper.resolve_vidalytics_poster(account_id, embed_id)
            if url and ripper.download_image(url, str(SCREENSHOTS_DIR / f'{job_id}_poster0.png')):
                return True
        except Exception:
            pass
    return False


def _headline_from_image(shot_path, api_key: str) -> dict:
    """Run the Claude vision headline extraction on a single image. Returns the
    parsed dict (may be empty). Raises on API failure so the caller can decide."""
    import base64
    import anthropic
    raw_bytes = shot_path.read_bytes()
    # Sniff the media type from magic bytes — poster images pulled from the
    # video CDN are usually JPEG, while our screenshots are PNG. Declaring the
    # wrong type makes the vision API 400 ("specified image/png … but the image").
    if raw_bytes[:3] == b'\xff\xd8\xff':
        media_type = 'image/jpeg'
    elif raw_bytes[:4] == b'RIFF' and raw_bytes[8:12] == b'WEBP':
        media_type = 'image/webp'
    elif raw_bytes[:6] in (b'GIF87a', b'GIF89a'):
        media_type = 'image/gif'
    else:
        media_type = 'image/png'
    img_b64 = base64.standard_b64encode(raw_bytes).decode('ascii')
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=HEADLINE_MODEL,
        max_tokens=500,
        messages=[{
            'role': 'user',
            'content': [
                {'type': 'image', 'source': {
                    'type': 'base64', 'media_type': media_type, 'data': img_b64,
                }},
                {'type': 'text', 'text': (
                    'This image is either the top of a financial promo landing '
                    'page or the pre-play thumbnail of its VSL video. '
                    'Extract the promo headline block, verbatim, in these parts:\n'
                    '- "eyebrow": the small pre-headline line ABOVE the headline '
                    '(e.g. "FORMER CIA ADVISOR RELEASES:").\n'
                    '- "headline": the single largest / most prominent line '
                    '(e.g. "THE AI BLACK PAPER").\n'
                    '- "subhead": the line directly under the headline '
                    '(e.g. "WARNS THE AI BUBBLE IS SET TO POP ON JULY 29TH AT 6:30PM").\n'
                    '- "subhead2": a secondary subhead or pull-quote, often below the '
                    'video (e.g. "The Dow Could Drop By 80% - Former CIA Advisor Jim '
                    'Rickards").\n'
                    'Which text to read: if there is promo text ABOVE the video, use '
                    'that; otherwise read the headline block from the video thumbnail '
                    'shown in the image.\n'
                    'Ignore only an obvious cookie/consent/privacy popup or an '
                    '"I AGREE" access box overlaying the page. Still read the main '
                    'promo hero — the largest heading and its sub-lines — including '
                    'confirmation-style headers (e.g. "MEETING CONFIRMED"). Return all '
                    'empty strings ONLY if the image is purely a bot/security '
                    'verification page with no marketing content.\n'
                    'Return ONLY compact JSON with keys eyebrow, headline, subhead, '
                    'subhead2. Use "" for any part that is not present. No other text.'
                )},
            ],
        }],
    )
    raw = ''.join(
        b.text for b in msg.content if getattr(b, 'type', '') == 'text'
    ).strip()
    if '{' in raw and '}' in raw:
        return json.loads(raw[raw.find('{'):raw.rfind('}') + 1])
    return {}


def _headline_from_text(hero_text: str, api_key: str) -> dict:
    """Structure raw hero text (captured from the promo page DOM) into the
    headline block. Used when the headline is real DOM text rather than an
    image. Returns the parsed dict (may be empty)."""
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=HEADLINE_MODEL,
        max_tokens=500,
        messages=[{
            'role': 'user',
            'content': [{'type': 'text', 'text': (
                'The following is the raw hero/top text scraped from a financial '
                'promo landing page (order roughly matches on-page order):\n\n'
                f'"""\n{hero_text}\n"""\n\n'
                'Extract the promo headline block, verbatim, in these parts:\n'
                '- "eyebrow": the small pre-headline line above the headline.\n'
                '- "headline": the single largest / most prominent line.\n'
                '- "subhead": the line directly under the headline.\n'
                '- "subhead2": a secondary subhead or pull-quote.\n'
                'Ignore nav links, buttons, disclaimers, boilerplate and cookie/'
                'consent popups. If there is no promo headline, return all empty '
                'strings.\n'
                'Return ONLY compact JSON with keys eyebrow, headline, subhead, '
                'subhead2. Use "" for any part not present. No other text.'
            )}],
        }],
    )
    raw = ''.join(
        b.text for b in msg.content if getattr(b, 'type', '') == 'text'
    ).strip()
    if '{' in raw and '}' in raw:
        return json.loads(raw[raw.find('{'):raw.rfind('}') + 1])
    return {}


def _extract_headline(job_id: str) -> None:
    """Read the promo headline block off the top-of-page screenshot via Claude
    vision and store it on the job. Best-effort — never raises, and no-ops when
    ANTHROPIC_API_KEY is unset so it can't break the pipeline.

    Captures the full headline hierarchy of a sales page:
      - eyebrow    : the small pre-headline line above the headline
                     (e.g. "FORMER CIA ADVISOR RELEASES:")
      - headline   : the largest, most prominent line (e.g. "THE AI BLACK PAPER")
      - subhead    : the line directly under the headline
      - subhead2   : a secondary subhead / pull-quote, often below the video
                     (e.g. "The Dow Could Drop By 80% — Former CIA Advisor Jim Rickards")

    Rule (matches the analyzer): the headline block is the text ABOVE the video;
    if there is none above the video, read it from the video thumbnail; otherwise
    ignore any text inside the thumbnail.
    """
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return
    # Candidate images, in priority order:
    #   1. page top-clip  — headline text ABOVE the video (normal promos)
    #   2. video poster   — headline baked into the thumbnail (Cloudflare-gated
    #                        promos we can't screenshot, e.g. Porter & Company)
    #   3. full-page      — last resort
    # Use the first image that yields a real headline; the page clip returns
    # empty for a Cloudflare "verify you are human" interstitial, so we fall
    # through to the poster automatically.
    # Priority: bookmarklet-captured page hero (the real headline on gated
    # promos) -> page top-clip -> video poster -> full page. The hero is what
    # the user actually sees; everything else is a fallback.
    # Fetch any bookmarklet-supplied hero image URLs server-side (cross-origin
    # heroes the browser canvas couldn't inline). These live on CDNs reachable
    # from the server even when the promo HTML page is Cloudflare-gated.
    _job = _get_job(job_id)
    _urls = (_job.get('hero_image_urls') if _job else None) or []
    if _urls:
        _idx = len(sorted(SCREENSHOTS_DIR.glob(f'{job_id}_hero*.png')))
        for _u in _urls[:6]:
            if _idx >= 4:
                break
            _dest = SCREENSHOTS_DIR / f'{job_id}_hero{_idx}.png'
            try:
                if ripper.download_image(_u, str(_dest)):
                    _idx += 1
            except Exception:
                pass
    hero_imgs = sorted(SCREENSHOTS_DIR.glob(f'{job_id}_hero*.png'))
    top = SCREENSHOTS_DIR / f'{job_id}_top.png'
    full = SCREENSHOTS_DIR / f'{job_id}.png'
    poster_imgs = sorted(SCREENSHOTS_DIR.glob(f'{job_id}_poster*.png'))
    if not hero_imgs and not poster_imgs:
        _ensure_poster(job_id)
        poster_imgs = sorted(SCREENSHOTS_DIR.glob(f'{job_id}_poster*.png'))
    # Order: real page hero (bookmarklet) → page top-clip → video frames
    # (earliest first, so a title card like "TRUMP'S NEW DOLLAR" at t=0 wins over
    # a later scene) → full page.
    candidates = [p for p in (*hero_imgs, top, *poster_imgs, full) if p.exists()]
    last_err = None
    for shot in candidates:
        try:
            data = _headline_from_image(shot, api_key)
        except Exception as exc:
            # One bad image (e.g. an API error) shouldn't abort the others.
            last_err = str(exc)[:200]
            continue
        fields = {
            'eyebrow': (data.get('eyebrow') or '').strip(),
            'headline': (data.get('headline') or '').strip(),
            'subhead': (data.get('subhead') or '').strip(),
            'subhead2': (data.get('subhead2') or '').strip(),
        }
        if any(fields.values()):
            _update_job(job_id, fields)
            return
    # No image yielded a headline — try the captured hero TEXT (headline may be
    # real DOM text rather than an image).
    job = _get_job(job_id)
    hero_text = (job.get('hero_text') if job else '') or ''
    if hero_text.strip():
        try:
            data = _headline_from_text(hero_text, api_key)
            fields = {
                'eyebrow': (data.get('eyebrow') or '').strip(),
                'headline': (data.get('headline') or '').strip(),
                'subhead': (data.get('subhead') or '').strip(),
                'subhead2': (data.get('subhead2') or '').strip(),
            }
            if any(fields.values()):
                _update_job(job_id, fields)
                return
        except Exception as exc:
            last_err = str(exc)[:200]
    # Nothing produced a headline — record the last error if there was one.
    if last_err:
        _update_job(job_id, {'headline_error': last_err})


def _run_pipeline(job_id: str, source_url: str) -> None:
    """
    Full rip pipeline: fetch → detect → download → screenshot → Rev submit.
    Runs in a background thread; updates manifest at each step.
    """
    try:
        # Step 1: fetch page (skip for direct player URLs — no useful metadata there)
        _update_job(job_id, {'pipeline_step': 'fetching_page'})
        if 'players.brightcove.net' in source_url:
            html, title, thumbnail = '', source_url, ''
        elif 'fast.wistia.com/medias/' in source_url or 'wistia.com/medias/' in source_url:
            # Direct Wistia media URL — extract ID from URL, skip page fetch
            import re as _re
            _m = _re.search(r'wistia\.com/medias/([a-z0-9]+)', source_url)
            html, title, thumbnail = '', source_url, ''
            if _m:
                _update_job(job_id, {
                    'platform': 'wistia',
                    'video_id': _m.group(1),
                    'pipeline_step': 'downloading_video',
                })
                dest_path = str(VIDEOS_DIR / f'{job_id}.mp4')
                ripper.download_video('wistia', _m.group(1), source_url, dest_path, {})
                _apply_video_metadata(job_id, source_url, 'wistia', _m.group(1), {})
                _update_job(job_id, {
                    'local_video': f'/data/videos/{job_id}.mp4',
                    'pipeline_step': 'taking_screenshot',
                })
                job = _get_job(job_id)
                # Screenshot the original promo page (page_url) for bookmarklet
                # jobs so the analyzer gets the headline/subheadline; fall back
                # to the source URL otherwise. Non-blocking on failure.
                _shot_url = (job.get('page_url') if job else None) or source_url
                _take_screenshot(job_id, _shot_url)
                _extract_headline(job_id)
                _update_job(job_id, {'pipeline_step': 'submitting_to_rev'})
                app_base = os.environ.get('APP_BASE_URL', 'https://vidripper.oxfordhub.app')
                video_url = f'{app_base}/api/jobs/{job_id}/video'
                now = datetime.now(timezone.utc).isoformat()
                order_id = rev_client.submit_job(video_url, metadata=job_id)
                _update_job(job_id, {
                    'rev_order_id': order_id,
                    'rev_status': 'in_progress',
                    'rev_submitted_at': now,
                    'pipeline_step': 'done',
                })
                return
        elif 'fast.vidalytics.com/embeds/' in source_url:
            # Direct Vidalytics embed URL — extract IDs from URL
            import re as _re
            _m = _re.search(r'fast\.vidalytics\.com/embeds/([A-Za-z0-9_-]+)/([A-Za-z0-9_-]+)', source_url)
            if _m:
                _acc, _vid = _m.group(1), _m.group(2)
                _update_job(job_id, {
                    'platform': 'vidalytics',
                    'video_id': _vid,
                    'pipeline_step': 'downloading_video',
                })
                dest_path = str(VIDEOS_DIR / f'{job_id}.mp4')
                ripper.download_video('vidalytics', _vid, source_url, dest_path, {'vidalytics_account_id': _acc})
                _apply_video_metadata(job_id, source_url, 'vidalytics', _vid, {'vidalytics_account_id': _acc})
                _update_job(job_id, {
                    'local_video': f'/data/videos/{job_id}.mp4',
                    'pipeline_step': 'taking_screenshot',
                })
                job = _get_job(job_id)
                # Screenshot the original promo page (page_url, set by the
                # bookmarklet) so the analyzer gets the headline/subheadline;
                # fall back to the embed URL otherwise. Non-blocking on failure.
                _shot_url = (job.get('page_url') if job else None) or source_url
                _take_screenshot(job_id, _shot_url)
                # The Vidalytics promo page is usually Cloudflare-gated (the
                # screenshot above just captures the bot check), and the headline
                # is baked into the video thumbnail — so grab the native poster
                # from the Vidalytics CDN for headline extraction to read.
                _ensure_poster(job_id, 'vidalytics', _acc, _vid)
                _extract_headline(job_id)
                _update_job(job_id, {'pipeline_step': 'submitting_to_rev'})
                app_base = os.environ.get('APP_BASE_URL', 'https://vidripper.oxfordhub.app')
                video_url = f'{app_base}/api/jobs/{job_id}/video'
                now = datetime.now(timezone.utc).isoformat()
                order_id = rev_client.submit_job(video_url, metadata=job_id)
                _update_job(job_id, {'rev_order_id': order_id, 'rev_status': 'in_progress', 'rev_submitted_at': now, 'pipeline_step': 'done'})
                return
            html, title, thumbnail = '', source_url, ''
        else:
            try:
                html, title, thumbnail = ripper.fetch_page(source_url)
            except Exception:
                # Plain-requests fetch can 403 on Cloudflare-protected promo
                # pages (e.g. Vidalytics landing pages). Fall through with empty
                # HTML so detection returns 'unknown' and the Playwright-rendered
                # retry below gets a real chance instead of failing the job.
                html, title, thumbnail = '', source_url, ''
        _update_job(job_id, {'title': title, 'thumbnail_url': thumbnail})

        # Domains known to work with yt-dlp's generic extractor — skip Playwright
        # rendering and pass source_url directly to yt-dlp.
        _GENERIC_EXTRACTOR_DOMAINS = (
            'foxbusiness.com',
            'foxnews.com',
        )

        # Step 2: detect platform — if unknown, retry with headless browser
        platform, video_id, extra = ripper.detect_platform(html, source_url)
        if platform == 'unknown':
            # Check if this is a domain that yt-dlp handles natively without rendering
            from urllib.parse import urlparse as _urlparse
            _host = _urlparse(source_url).netloc.lower().lstrip('www.')
            _skip_playwright = any(_host.endswith(d) for d in _GENERIC_EXTRACTOR_DOMAINS)

            if not _skip_playwright:
                _update_job(job_id, {'pipeline_step': 'fetching_page_rendered'})
                html, title_r, thumbnail_r, bc_attrs = ripper.fetch_page_rendered(source_url)
                if title_r and title_r != source_url:
                    _update_job(job_id, {'title': title_r, 'thumbnail_url': thumbnail_r})
                # If Playwright found BrightCove attrs directly in the DOM, use them
                if bc_attrs and bc_attrs.get('video_id') and bc_attrs.get('account_id'):
                    platform = 'brightcove'
                    video_id = bc_attrs['video_id']
                    extra = {
                        'bc_account_id': bc_attrs['account_id'],
                        'bc_player_id': bc_attrs.get('player_id', ''),
                    }
                else:
                    platform, video_id, extra = ripper.detect_platform(html, source_url)
        _update_job(job_id, {
            'platform': platform,
            'video_id': video_id,
            'pipeline_step': 'downloading_video',
        })

        # Step 3: download
        dest_path = str(VIDEOS_DIR / f'{job_id}.mp4')
        ripper.download_video(platform, video_id, source_url, dest_path, extra)
        # Capture the real yt-dlp video title + thumbnail fallback.
        current = _get_job(job_id)
        _page_title = (current.get('title') if current else '') or title
        _apply_video_metadata(job_id, source_url, platform, video_id, extra, _page_title)
        _update_job(job_id, {
            'local_video': f'/data/videos/{job_id}.mp4',
            'pipeline_step': 'taking_screenshot',
        })

        # Step 4: full-page screenshot (non-blocking on failure).
        # Prefer the original promo page (page_url, set for bookmarklet jobs) so
        # the analyzer can read the headline/subheadline; fall back to the
        # pasted URL. Uploaded cookies let gated pages render authenticated.
        job = _get_job(job_id)
        _shot_url = (job.get('page_url') if job else None) or source_url
        # Skip screenshots for generic-extractor domains (e.g. Fox Business) —
        # those are news clips, not promos, so there's no headline to capture.
        from urllib.parse import urlparse as _up_shot
        _shot_host = _up_shot(_shot_url).netloc.lower().lstrip('www.')
        if not any(_shot_host.endswith(d) for d in _GENERIC_EXTRACTOR_DOMAINS):
            _take_screenshot(job_id, _shot_url)
            _extract_headline(job_id)
        _update_job(job_id, {'pipeline_step': 'submitting_to_rev'})

        # Step 5: Rev AI — submit video URL directly, no upload needed
        app_base = os.environ.get('APP_BASE_URL', 'https://vidripper.oxfordhub.app')
        video_url = f'{app_base}/api/jobs/{job_id}/video'
        now = datetime.now(timezone.utc).isoformat()
        order_id = rev_client.submit_job(video_url, metadata=job_id)
        _update_job(job_id, {
            'rev_order_id': order_id,
            'rev_status': 'in_progress',
            'rev_submitted_at': now,
            'pipeline_step': 'done',
        })

    except Exception as exc:
        _update_job(job_id, {
            'rev_status': 'failed',
            'pipeline_step': 'failed',
            'error': str(exc)[:500],
        })


# ── routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory(str(FRONTEND_DIR), 'index.html')


@app.route('/favicon.ico')
def favicon():
    return send_from_directory(str(FRONTEND_DIR), 'favicon.ico', mimetype='image/x-icon')


@app.route('/icon.png')
def icon_png():
    return send_from_directory(str(FRONTEND_DIR), 'icon.png', mimetype='image/png')


@app.route('/api/ping')
def ping():
    return jsonify({'status': 'ok', 'service': 'vidripper'})


@app.route('/api/diag')
def diag():
    """Read-only diagnostics to detect data-durability problems: which instance
    served this request, whether DATA_DIR is a real persistent mount, and the
    manifest size. Hit repeatedly — if hostname/deployment vary or the manifest
    count jumps, storage is per-instance (not a shared volume) and jobs/
    screenshots will randomly go missing."""
    import socket
    try:
        _count = len(_load_manifest())
    except Exception:
        _count = -1
    return jsonify({
        'hostname': socket.gethostname(),
        'data_dir': str(DATA_DIR),
        # Check the ACTUAL DATA_DIR (the volume is mounted at /app/data, not
        # /data), so this reflects real persistence instead of a false negative.
        'data_is_mount': os.path.ismount(str(DATA_DIR)),
        'manifest_count': _count,
        'screenshot_files': len(list(SCREENSHOTS_DIR.glob('*.png'))),
        'residential_proxy_configured': bool(ripper.cnn_proxy_url()),
        'railway': {
            k: os.environ.get(k) for k in (
                'RAILWAY_REPLICA_ID', 'RAILWAY_DEPLOYMENT_ID',
                'RAILWAY_SERVICE_NAME', 'RAILWAY_PROJECT_NAME',
                'RAILWAY_REPLICA_REGION',
            )
        },
    })


@app.route('/api/rip', methods=['POST', 'OPTIONS'])
def rip():
    if request.method == 'OPTIONS':
        return ('', 204)  # CORS preflight
    data = request.get_json(silent=True) or {}
    url = (data.get('url') or '').strip()
    if not url:
        return jsonify({'error': 'url is required'}), 400
    if not url.startswith(('http://', 'https://')):
        return jsonify({'error': 'url must start with http:// or https://'}), 400

    # ── Duplicate check ──────────────────────────────────────────────────────
    # Skip it when the caller supplies freshly-captured hero data (a bookmarklet
    # re-run): the whole point of re-running is to re-capture the hero/headline,
    # so returning the stale job would defeat it. Manual pastes (no hero data)
    # still dedupe.
    _has_hero = bool(
        data.get('hero_image_urls') or data.get('hero_images')
        or data.get('hero_image') or (data.get('hero_text') or '').strip()
    )
    if not _has_hero:
        existing = _find_duplicate(url)
        if existing:
            return jsonify({'duplicate': True, 'existing_job': existing}), 200

    job_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    page_url = (data.get('page_url') or '').strip() or None

    # ── Hero capture from the bookmarklet ────────────────────────────────────
    # On Cloudflare-gated promos the headline lives in the page's hero (image
    # and/or heading text), which only renders in the user's already-cleared
    # browser. The bookmarklet grabs it there and posts it; we persist the hero
    # image(s) for vision extraction and the text as a fallback.
    hero_text = (data.get('hero_text') or '').strip()[:2000]
    _heroes = data.get('hero_images') or []
    if isinstance(data.get('hero_image'), str):
        _heroes = [data['hero_image'], *(_heroes if isinstance(_heroes, list) else [])]
    _saved = 0
    if isinstance(_heroes, list):
        import base64 as _b64
        for du in _heroes[:3]:
            if not isinstance(du, str) or ',' not in du or not du.startswith('data:'):
                continue
            try:
                raw = _b64.b64decode(du.split(',', 1)[1])
            except Exception:
                continue
            if not raw or len(raw) > 8 * 1024 * 1024:
                continue
            (SCREENSHOTS_DIR / f'{job_id}_hero{_saved}.png').write_bytes(raw)
            _saved += 1
    # Candidate hero image URLs (largest-first). Cross-origin hero images taint
    # the browser canvas so the bookmarklet can't inline them; instead it sends
    # the URLs and we fetch them server-side from the CDN (which, unlike the
    # gated HTML page, is reachable) during headline extraction.
    hero_image_urls = [
        u for u in (data.get('hero_image_urls') or [])
        if isinstance(u, str) and u.startswith('http')
    ][:6]

    job = {
        'id': job_id,
        'source_url': url,
        'page_url': page_url,
        'platform': 'unknown',
        'video_id': '',
        'title': '',
        'thumbnail_url': '',
        'local_video': '',
        'rev_order_id': '',
        'rev_status': 'pending',
        'rev_submitted_at': '',
        'rev_transcript_url': '',
        'transcript_text': '',
        'has_screenshot': False,
        'hero_text': hero_text,
        'hero_image_urls': hero_image_urls,
        'eyebrow': '',
        'headline': '',
        'subhead': '',
        'subhead2': '',
        'analyze_status': '',
        'analyze_error': '',
        'drive_file_id': '',
        'drive_url': '',
        'archived_at': '',
        'archive_error': '',
        'pipeline_step': 'queued',
        'error': '',
        'created_at': now,
    }
    _append_job(job)

    thread = threading.Thread(target=_run_pipeline, args=(job_id, url), daemon=True)
    thread.start()

    return jsonify(job), 202


@app.route('/api/jobs')
def list_jobs():
    jobs = _load_manifest()
    # Most recent first
    jobs.sort(key=lambda j: j.get('created_at', ''), reverse=True)
    return jsonify(jobs)


@app.route('/api/jobs/<job_id>')
def get_job(job_id):
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'id': job_id}), 404
    return jsonify(job)


@app.route('/api/jobs/<job_id>/transcript')
def get_transcript(job_id):
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'id': job_id}), 404

    # Return cached transcript if present
    if job.get('transcript_text'):
        return jsonify({'transcript': job['transcript_text'], 'source': 'cache'})

    order_id = job.get('rev_order_id')
    if not order_id:
        return jsonify({'error': 'No Rev order associated with this job'}), 400

    try:
        status = rev_client.get_job_status(order_id)
        _update_job(job_id, {'rev_status': status})

        if status == 'failed':
            return jsonify({'error': 'Transcription failed in Rev AI', 'status': 'failed'}), 500
        if status != 'complete':
            return jsonify({'status': status, 'message': 'Transcript not ready yet'}), 202

        text = rev_client.get_transcript(order_id)
        _update_job(job_id, {'transcript_text': text, 'rev_status': 'complete'})
        return jsonify({'transcript': text, 'source': 'revai'})

    except Exception as exc:
        return jsonify({'error': str(exc)}), 500


@app.route('/api/jobs/<job_id>/transcript.docx')
def get_transcript_docx(job_id):
    """Download transcript as a formatted Word document."""
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'id': job_id}), 404

    text = job.get('transcript_text', '')
    if not text:
        return jsonify({'error': 'Transcript not available yet'}), 400

    try:
        docx_bytes = _build_docx(
            job.get('title', ''), text,
            job.get('eyebrow', ''), job.get('headline', ''),
            job.get('subhead', ''), job.get('subhead2', ''),
        )
    except Exception as exc:
        return jsonify({'error': str(exc)}), 500

    safe_title = ''.join(c for c in (job.get('title') or job_id)[:40] if c.isalnum() or c in ' -_').strip() or job_id
    filename = f'{safe_title}.docx'

    return Response(
        docx_bytes,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(docx_bytes)),
        },
    )


@app.route('/api/jobs/<job_id>/screenshot')
def get_screenshot(job_id):
    """Download the page screenshot PNG."""
    screenshot_path = SCREENSHOTS_DIR / f'{job_id}.png'
    if not screenshot_path.exists():
        return jsonify({'error': 'Screenshot not found', 'id': job_id}), 404
    return send_from_directory(str(SCREENSHOTS_DIR), f'{job_id}.png', mimetype='image/png')


@app.route('/api/jobs/<job_id>/video')
def serve_video(job_id):
    # Local file present → serve it directly (fast path, supports Range natively).
    video_path = VIDEOS_DIR / f'{job_id}.mp4'
    if video_path.exists():
        return send_from_directory(str(VIDEOS_DIR), f'{job_id}.mp4', mimetype='video/mp4')
    # Archived → stream from Google Drive, forwarding Range so the player seeks.
    job = _get_job(job_id)
    if job and job.get('drive_file_id') and gdrive.is_configured():
        try:
            up = gdrive.open_stream(job['drive_file_id'], request.headers.get('Range'))
        except Exception as exc:
            return jsonify({'error': f'drive stream failed: {exc}'}), 502
        if up.status_code >= 400:
            up.close()
            return jsonify({'error': 'archived video unavailable',
                            'drive_status': up.status_code}), 502
        passthrough = ('content-length', 'content-range', 'content-type', 'accept-ranges')
        headers = {k: v for k, v in up.headers.items() if k.lower() in passthrough}
        headers.setdefault('Accept-Ranges', 'bytes')
        return Response(
            up.iter_content(chunk_size=262144),
            status=up.status_code,
            headers=headers,
            mimetype='video/mp4',
        )
    return jsonify({'error': 'Job not found', 'id': job_id}), 404


@app.route('/api/admin/archive', methods=['POST'])
def admin_archive():
    """Manually trigger the archival sweep. Optional ?days=N overrides the
    30-day threshold (use ?days=0 to archive everything now, for testing)."""
    days = request.args.get('days', type=int)
    return jsonify(_archive_old_videos(days=days))


@app.route('/api/jobs/<job_id>', methods=['DELETE'])
def delete_job(job_id):
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'id': job_id}), 404

    # Remove video file if present
    video_path = VIDEOS_DIR / f'{job_id}.mp4'
    if video_path.exists():
        try:
            video_path.unlink()
        except OSError:
            pass

    # Remove the archived Drive copy too, if this job was offloaded.
    if job.get('drive_file_id') and gdrive.is_configured():
        try:
            gdrive.delete_file(job['drive_file_id'])
        except Exception:
            pass

    # Remove screenshots (full page + top clip + poster + captured heroes)
    _shots = [f'{job_id}.png', f'{job_id}_top.png']
    _shots += [p.name for p in SCREENSHOTS_DIR.glob(f'{job_id}_poster*.png')]
    _shots += [p.name for p in SCREENSHOTS_DIR.glob(f'{job_id}_hero*.png')]
    for _shot in _shots:
        _sp = SCREENSHOTS_DIR / _shot
        if _sp.exists():
            try:
                _sp.unlink()
            except OSError:
                pass

    _delete_job(job_id)
    return jsonify({'deleted': True, 'id': job_id})



def _run_analysis(job_id: str, cookie: str) -> None:
    """Background worker: build the transcript .docx (+ screenshot), POST it to
    the Promo Analyzer, and store the returned review id on the job.

    Runs in a thread because a full analysis can exceed the platform's HTTP
    request timeout — a synchronous proxy hit Railway's edge limit and returned
    "upstream error". The frontend polls the job for promo_review_id /
    analyze_status instead. Best-effort — never raises.
    """
    import re as _re
    import requests as _requests
    try:
        job = _get_job(job_id)
        if not job:
            return
        text = job.get('transcript_text', '')
        title = job.get('title', '') or job_id
        docx_bytes = _build_docx(
            title, text,
            job.get('eyebrow', ''), job.get('headline', ''),
            job.get('subhead', ''), job.get('subhead2', ''),
        )
        safe_title = ''.join(c for c in title[:40] if c.isalnum() or c in ' -_').strip() or job_id
        filename = f'{safe_title}.docx'

        # Send only the transcript .docx. The extracted headline/subheadline is
        # already rendered at the top of that doc, so the analyzer does not need
        # the screenshot image — keeping the payload small avoids extra upload
        # latency on the analyze call.
        files = {
            'file': (filename, docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        }

        # Forward the OxfordHub session cookie so the analyzer's auth resolves the
        # signed-in user; fall back to the server-to-server token.
        fwd_headers = {}
        if cookie:
            fwd_headers['Cookie'] = cookie
        _hub_token = os.environ.get('HUB_API_TOKEN')
        if _hub_token:
            fwd_headers['x-hub-token'] = _hub_token

        analyzer_url = 'https://analyzer.oxfordhub.app/api/analyze'
        resp = _requests.post(
            analyzer_url, files=files, headers=fwd_headers, timeout=(15, 600),
        )
        full_text = resp.text
        review_id = None
        meta_match = _re.search(r'\[META\](.*?)\[/META\]', full_text, _re.DOTALL)
        if meta_match:
            try:
                review_id = json.loads(meta_match.group(1)).get('reviewId')
            except Exception:
                pass
        if review_id:
            _update_job(job_id, {
                'promo_review_id': review_id,
                'analyze_status': 'done',
                'analyze_error': '',
            })
        else:
            _update_job(job_id, {
                'analyze_status': 'error',
                'analyze_error': f'Analyzer returned no review id: {full_text[:180]}',
            })
    except Exception as exc:
        _update_job(job_id, {'analyze_status': 'error', 'analyze_error': str(exc)[:200]})


@app.route('/api/jobs/<job_id>/analyze-proxy', methods=['POST'])
def analyze_proxy(job_id):
    """Kick off analysis in the background and return immediately (202).

    The frontend polls the job for promo_review_id / analyze_status. Async
    because a full analysis can outlast the platform HTTP request timeout.
    """
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    if not job.get('transcript_text'):
        return jsonify({'error': 'Transcript not available yet'}), 400
    if job.get('promo_review_id'):
        return jsonify({'status': 'done', 'review_id': job['promo_review_id']}), 200
    if job.get('analyze_status') == 'analyzing':
        return jsonify({'status': 'analyzing'}), 202

    # Capture the hub cookie now (request context) to forward from the thread.
    cookie = request.headers.get('Cookie', '')
    _update_job(job_id, {'analyze_status': 'analyzing', 'analyze_error': ''})
    threading.Thread(target=_run_analysis, args=(job_id, cookie), daemon=True).start()
    return jsonify({'status': 'analyzing'}), 202

@app.route('/api/jobs/<job_id>/set-review', methods=['POST'])
def set_review(job_id):
    """Persist the Promo Analyzer review ID on a job."""
    job = _get_job(job_id)
    if not job:
        return jsonify({'error': 'Job not found', 'id': job_id}), 404
    data = request.get_json(silent=True) or {}
    review_id = (data.get('promo_review_id') or '').strip()
    if not review_id:
        return jsonify({'error': 'promo_review_id is required'}), 400
    updated = _update_job(job_id, {'promo_review_id': review_id})
    return jsonify({'ok': True, 'promo_review_id': review_id})


COOKIES_DIR = DATA_DIR / 'cookies'


@app.route('/api/admin/upload-cookies', methods=['POST'])
def upload_cookies():
    """Upload a Netscape cookies.txt file for a domain."""
    f = request.files.get('file')
    domain = (request.form.get('domain') or '').strip()
    if not f or not domain:
        return jsonify({'error': 'file and domain are required'}), 400
    COOKIES_DIR.mkdir(parents=True, exist_ok=True)
    dest = COOKIES_DIR / f'{domain}.txt'
    f.save(str(dest))
    return jsonify({'saved': str(dest), 'domain': domain})


if __name__ == '__main__':
    app.run(debug=True, port=5001)
